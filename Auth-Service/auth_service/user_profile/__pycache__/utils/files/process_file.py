from django.core.files import File
from typing import Optional, Dict, Any
import os
import hashlib
import logging
from magic import Magic
from pathlib import Path
from django.core.files.storage import default_storage
from PyPDF2 import PdfReader
import docx
from pptx import Presentation
import magic


logger = logging.getLogger('utils')

class FileProcessor:
    """
    Processes uploaded files to extract metadata, content, and MIME type.
    Follows Single Responsibility Principle (SRP) by separating file operations.
    """

    def __init__(self, file_obj: File):
        self.file = file_obj  # Django File object
        self._validate_file()

    def _validate_file(self) -> None:
        """Ensure the file is accessible and valid."""
        if not hasattr(self.file, 'name'):
            raise ValueError("File name Not Found")
        if not self.file:
            raise ValueError("File object cannot be None")
        if not hasattr(self.file, 'read'):
            raise TypeError("Expected Django File-like object")
        

    def get_metadata(self) -> Dict[str, Any]:
        """
        Extracts core file metadata (size, checksum, MIME type).
        Returns: {
            'size': int,
            'checksum': str,
            'mime_type': str,
            'extension': str
        }
        """
        try:
            return {
                'size': self._get_file_size(),
                'checksum': self._calculate_checksum(),
                'mime_type': self._detect_mime_type(),
                'extension': self._get_extension()
            }
        except Exception as e:
            logger.error(f"Metadata extraction failed: {e} to: \n {Path(__file__).resolve()}")
            raise NotImplementedError("This is abstract Method can not used to create object.")                                                                                                                                                                                                                                               

    def _get_file_size(self) -> int:
        """Returns file size in bytes."""
        with default_storage.open(self.file.name, 'rb') as file_object:
            file_object.seek(0, os.SEEK_END)
            size = file_object.tell()
            file_object.seek(0)  # Reset pointer
            return size

    def _calculate_checksum(self, algorithm: str = 'sha256') -> str:
        """Generates a checksum for file integrity verification."""
        with default_storage.open(self.file.name, 'rb') as file_object:
            hash_func = hashlib.new(algorithm)
            for chunk in file_object.chunks(4096):
                hash_func.update(chunk)
            return hash_func.hexdigest()

    def _detect_mime_type(self) -> str:
        """Detects MIME type using libmagic with fallback to extensions."""
        try:
            with default_storage.open(self.file.name, 'rb') as file_object: # 'rb' for binary read
                mime = Magic(mime=True)
                chunk = file_object.read(1024)
                return mime.from_buffer(chunk)
        except Exception:
            # Fallback to extension-based detection
            ext = self._get_extension().lower()
            return {
                '.pdf': 'application/pdf',
                '.jpg': 'image/jpeg',
                '.png': 'image/png',
                # Add more mappings as needed
            }.get(ext, 'application/octet-stream')

    def _get_extension(self) -> str:
        """Extracts file extension from the name."""
        return Path(self.file.name).suffix.lower()

    def extract_text(self) -> Optional[str]:
        """Placeholder for text extraction (override for specific formats)."""
        # Implement with PyPDF2, python-docx, etc. as needed
        raise NotImplementedError("The Implementations of extract_text() is not yet done")
    
    
class PDFProcessor(FileProcessor):
    """Handles PDF-specific operations."""
    def extract_text(self) -> Optional[str]:
        try:
            from PyPDF2 import PdfReader
            with default_storage.open(self.file.name, 'rb') as file_object:
                reader = PdfReader(file_object)
                return "\n".join(page.extract_text() for page in reader.pages)
    
        except ImportError:
            logger.warning("PyPDF2 not installed for PDF extraction")
            return None
        
    def get_page_number(self):
        try:
            from PyPDF2 import PdfReader
            with default_storage.open(self.file.name, 'rb') as file_object:
                reader = PdfReader(file_object)
                return len(reader.pages)
        except ImportError:
            logger.warning("PyPDF2 not installed for PDF extraction")
            return 0

class DocumentProcessor(FileProcessor):
    """Handles document-specific operations."""
    def extract_text(self) -> Optional[str]:
        try:
            from docx import Document
            with default_storage.open(self.file.name, 'rb') as file_object:
                document = Document(file_object)
                return "\n".join(paragraph.text for paragraph in document.paragraphs)  
        except ImportError:
            logger.warning("python-docx not installed for document extraction")
            return None
    
    def get_page_number(self):
        """Counts paragraphs in a .docx file."""
        from docx import Document
        try:
            with default_storage.open(self.file.name, 'rb') as file_object:
                doc = Document(file_object)
                return len(doc.paragraphs)
        except Exception:
            return 0
        
class TextProcessor(FileProcessor):
    """Handles text-specific operations."""
    def extract_text(self) -> Optional[str]:
        with default_storage.open(self.file.name, 'r', encoding='utf-8') as f:
            return "\n".join(i.strip() for i in f.readlines())
        
    # def get_page_number(self):  
    #     with default_storage.open(self.file.name, 'r', encoding='utf-8') as f:
    #         return len(f.readlines())
        

class ImageProcessor(FileProcessor):
    """Handles image-specific metadata."""
    def get_metadata(self) -> Dict[str, Any]:
        metadata = super().get_metadata()
        metadata.update({'dimensions': self._get_image_dimensions()})
        return metadata

    def _get_image_dimensions(self) -> Optional[tuple]:
        """Uses PIL to extract image dimensions."""
        try:
            from PIL import Image

            with default_storage.open(self.file.name, 'rb') as file_object:
                image_bytes = file_object.read()
                with Image.open(io.BytesIO(image_bytes)) as img:
                    return img.size
        except ImportError:
            logger.warning("Pillow not installed for image processing")
            return None


class DocumentPageCounter:
    """
    Utility class for counting pages/slides/lines in various document types.

    Supported file types:
    - PDF (.pdf): Counts number of pages.
    - Word (.docx): Counts number of paragraphs (not actual pages).
    - PowerPoint (.pptx): Counts number of slides.
    - Plain text (.txt): Counts number of lines.
    """

    file_name: str
    _mime_detector: magic.Magic
    mime_type: Optional[str]

    def __init__(self, file: File) -> None:
        """
        Initializes the counter with a Django File object.
        
        :param file: Django File object
        """
        self.file_name: str = file.name
        self._mime_detector: magic.Magic = magic.Magic(mime=True)
        self.mime_type: Optional[str] = self._get_mime_type()

    def _get_mime_type(self) -> Optional[str]:
        """
        Detects the MIME type of the file using the magic library.

        :return: MIME type string or None on error
        """
        try:
            return self._mime_detector.from_file(default_storage.path(self.file_name))
        except Exception as e:
            logger.error(f"Error detecting MIME type: {e}")
            return None

    def count_pages(self) -> Optional[int]:
        """
        Counts pages/slides/lines depending on file type.

        :return: Integer count or None if unsupported/error
        """
        if not self.mime_type:
            return None

        if "pdf" in self.mime_type:
            return self._count_pdf_pages()
        elif "wordprocessingml.document" in self.mime_type:  # .docx
            return self._count_docx_paragraphs()
        elif "presentationml.presentation" in self.mime_type:  # .pptx
            return self._count_pptx_slides()
        elif "text/plain" in self.mime_type:
            return self._count_txt_lines()
        else:
            logger.warning(f"Unsupported MIME type: {self.mime_type}")
            return None

    def _count_pdf_pages(self) -> Optional[int]:
        """
        Counts the number of pages in a PDF file.

        :return: Page count or None on error
        """
        try:
            with default_storage.open(self.file_name, 'rb') as file_object:
                reader: PdfReader = PdfReader(file_object)
                return len(reader.pages)
        except Exception as e:
            logger.error(f"Error counting PDF pages: {e}")
            return None

    def _count_docx_paragraphs(self) -> Optional[int]:
        """
        Counts the number of paragraphs in a .docx file.
        (Note: Not an accurate page count.)

        :return: Paragraph count or None on error
        """
        try:
            with default_storage.open(self.file_name, 'rb') as file_object:
                doc: docx.Document = docx.Document(file_object)
                return len(doc.paragraphs)
        except Exception as e:
            logger.error(f"Error counting paragraphs in .docx: {e}")
            return None

    def _count_pptx_slides(self) -> Optional[int]:
        """
        Counts the number of slides in a .pptx file.

        :return: Slide count or None on error
        """
        try:
            with default_storage.open(self.file_name, 'rb') as file_object:
                presentation: Presentation = Presentation(file_object)
                return len(presentation.slides)
        except Exception as e:
            logger.error(f"Error counting slides in .pptx: {e}")
            return None

    def _count_txt_lines(self) -> Optional[int]:
        """
        Counts the number of lines in a plain text file.

        :return: Line count or None on error
        """
        try:
            with default_storage.open(self.file_name, 'r', encoding='utf-8') as file_object:
                return sum(1 for _ in file_object)
        except Exception as e:
            logger.error(f"Error counting lines in text file: {e}")
            return None
